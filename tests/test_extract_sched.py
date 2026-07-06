"""Schedule-genre extraction: unit rules + the artemis 3-day oracle.

The oracle idea: the doc's day tables are its own worked example. We extract
RULES (not pinned times), so (a) the plan must solve OPTIMAL on the 3-day
horizon, and (b) the doc's own timetable must pass the extracted rules —
except where the document contradicts ITSELF, which the self-check must
surface (day-2 sleep runs 22:00 + 8h30m = 06:30, but day 3 starts at 06:20).
"""
import io

import pytest
from docx import Document

from conftest import TESTDATA, fail_ask
import extract_sched as xs
from extract import detect_genre, extract_document
from models import Scenario
from solver import solve

ARTEMIS = TESTDATA / "artemis_3day_schedule.docx"


@pytest.fixture(scope="session")
def artemis_blocks():
    assert ARTEMIS.exists(), (
        "testdata/artemis_3day_schedule.docx is missing — it is the schedule-genre "
        "target document (untracked at first; commit it with the extractor)."
    )
    from ingest import extract_blocks
    return extract_blocks(ARTEMIS.read_bytes())["blocks"]


@pytest.fixture(scope="session")
def artemis(artemis_blocks):
    return extract_document(artemis_blocks, ask=fail_ask)  # deterministic — no LLM


# ---- small parsers ---------------------------------------------------------
def test_dur_to_min():
    assert xs.dur_to_min("45m") == 45
    assert xs.dur_to_min("2h") == 120
    assert xs.dur_to_min("8h15m") == 495
    assert xs.dur_to_min("1h10m") == 70
    assert xs.dur_to_min("") is None
    assert xs.dur_to_min("soon") is None


def test_hhmm_to_min():
    assert xs.hhmm_to_min("06:30") == 390
    assert xs.hhmm_to_min("22:15") == 1335
    assert xs.hhmm_to_min("6h") is None


def test_base_key_strips_day_variants():
    assert xs._base_key("Adaptation D1") == "adaptation"
    assert xs._base_key("Work C") == "work_c"
    assert xs._base_key("Prep/Light Duties") == "prep_light_duties"


# ---- genre detection -------------------------------------------------------
def test_genre_detection(artemis_blocks, sample_blocks):
    assert detect_genre(artemis_blocks) == "schedule"
    assert detect_genre(sample_blocks["blocks"]) == "spec"
    assert detect_genre([]) == "spec"  # nothing recognizable -> the spec path reports it


# ---- a tiny synthetic schedule doc (rules exercised in isolation) ----------
def mini_doc(bullets, days):
    """Build an in-memory schedule .docx: rule bullets + one table per day."""
    doc = Document()
    doc.add_paragraph("Constraints").runs  # plain preamble
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    for i, rows in enumerate(days):
        p = doc.add_paragraph(f"Day {i + 1}")
        p.runs[0].bold = True
        t = doc.add_table(rows=1 + len(rows), cols=5)
        for c, h in enumerate(["Start", "End", "Duration", "Activity", "Notes"]):
            t.rows[0].cells[c].text = h
        for r, row in enumerate(rows, start=1):
            for c, v in enumerate(row):
                t.rows[r].cells[c].text = v
    buf = io.BytesIO()
    doc.save(buf)
    from ingest import extract_blocks
    return extract_blocks(buf.getvalue())["blocks"]


def test_mini_roster_recurring_vs_variant():
    blocks = mini_doc([], [
        [["08:00", "09:00", "1h", "Checkout", ""], ["09:10", "10:40", "90m", "Cal D1", ""]],
        [["08:00", "09:00", "1h", "Checkout", ""], ["09:10", "11:10", "2h", "Cal D2", ""]],
    ])
    res = xs.extract_schedule(blocks)
    acts = {a["id"]: a for a in res["scenario"]["activities"]}
    # Same duration every day -> one recurring activity over all days.
    assert acts["checkout"]["recurs_daily"] is True
    assert acts["checkout"]["days"] == "all"
    # A 'D1'/'D2' name -> per-day variants pinned to their day.
    assert acts["cal_d1"]["days"] == [0] and acts["cal_d1"]["duration"] == 90
    assert acts["cal_d2"]["days"] == [1] and acts["cal_d2"]["duration"] == 120


def test_mini_duration_varying_name_splits():
    blocks = mini_doc([], [
        [["08:00", "09:00", "1h", "Ops", ""]],
        [["08:00", "10:00", "2h", "Ops", ""]],
    ])
    res = xs.extract_schedule(blocks)
    ids = sorted(a["id"] for a in res["scenario"]["activities"])
    assert ids == ["ops_d1", "ops_d2"]  # same name, different duration -> split per day


def test_mini_unmodeled_bullet_is_warned_not_dropped():
    blocks = mini_doc(["Thermal soak must respect beta angle limits"], [
        [["08:00", "09:00", "1h", "Ops", ""]],
    ])
    res = xs.extract_schedule(blocks)
    ledger = res["coverage"]["bullets"]
    assert len(ledger) == 1 and ledger[0]["status"] == "unmodeled"
    assert any("Un-modeled rule" in w for w in res["warnings"])
    assert res["coverage"]["n_constraints"] == 0  # never a fabricated constraint


def test_mini_rule_duration_beats_table_and_warns():
    blocks = mini_doc(["Nap: 30m contiguous"], [
        [["13:00", "13:45", "45m", "Nap", ""]],
    ])
    res = xs.extract_schedule(blocks)
    nap = next(a for a in res["scenario"]["activities"] if a["id"] == "nap")
    assert nap["duration"] == 30
    assert any("duration differs" in w for w in res["warnings"])


def test_mini_unreadable_duration_row_degrades_not_crashes():
    # Regression: a group with no readable duration is skipped by finalize_roster;
    # the coverage rows and any bullet rules referencing it must skip too (KeyError
    # here used to 500 the whole /extract).
    blocks = mini_doc([">=10m between activities"], [
        [["TBD", "", "", "Mystery Task", ""], ["08:00", "09:00", "1h", "Ops", ""]],
    ])
    res = xs.extract_schedule(blocks)
    assert any("no readable duration" in w for w in res["warnings"])
    ids = {a["id"] for a in res["scenario"]["activities"]}
    assert ids == {"ops"}
    rows = res["coverage"]["rows"]
    assert next(r for r in rows if "Mystery" in r["source"])["activity"] is None


def test_mini_day_zero_numbering_is_normalized():
    blocks = mini_doc([], [
        [["08:00", "09:00", "1h", "Ops", ""]],
        [["08:00", "09:00", "1h", "Ops", ""]],
    ])
    # rewrite the headings to a 0-based doc ("Day 1"/"Day 2" -> can't reuse mini_doc
    # directly), so build one by hand with Day 0 / Day 1 headings
    import io as _io
    from docx import Document as _D
    doc = _D()
    for n in (0, 1):
        p = doc.add_paragraph(f"Day {n}")
        p.runs[0].bold = True
        t = doc.add_table(rows=2, cols=5)
        for c, h in enumerate(["Start", "End", "Duration", "Activity", "Notes"]):
            t.rows[0].cells[c].text = h
        for c, v in enumerate(["08:00", "09:00", "1h", "Ops", ""]):
            t.rows[1].cells[c].text = v
    buf = _io.BytesIO()
    doc.save(buf)
    from ingest import extract_blocks
    res = xs.extract_schedule(extract_blocks(buf.getvalue())["blocks"])
    ops = next(a for a in res["scenario"]["activities"] if a["id"] == "ops")
    assert ops["days"] == "all" and res["coverage"]["n_days"] == 2  # 0-based doc still maps to days 0..1


def test_mini_zero_length_row_is_skipped_not_24h():
    blocks = mini_doc([], [
        [["08:00", "08:00", "", "Milestone", ""], ["09:00", "10:00", "1h", "Ops", ""]],
    ])
    res = xs.extract_schedule(blocks)
    assert {a["id"] for a in res["scenario"]["activities"]} == {"ops"}
    assert any("no readable duration" in w for w in res["warnings"])


def test_mini_duplicate_name_in_one_day_keeps_one_id():
    # Same name twice in one day + varying across days -> the per-day split must
    # still emit ONE id per day (duplicate ids would fail Scenario validation).
    blocks = mini_doc([], [
        [["08:00", "09:00", "1h", "Ops", ""], ["10:00", "10:30", "30m", "Ops", ""]],
        [["08:00", "10:00", "2h", "Ops", ""]],
    ])
    res = xs.extract_schedule(blocks)
    ids = sorted(a["id"] for a in res["scenario"]["activities"])
    assert ids == ["ops_d1", "ops_d2"]
    d1 = next(a for a in res["scenario"]["activities"] if a["id"] == "ops_d1")
    assert d1["duration"] == 30  # the last instance that day, matching the warning
    assert any("using the last one" in w for w in res["warnings"])


# ---- the artemis oracle ----------------------------------------------------
def test_artemis_coverage(artemis):
    cov = artemis["coverage"]
    assert cov["genre"] == "schedule"
    assert cov["n_tables"] == 3 and cov["n_days"] == 3 and cov["n_rows"] == 39
    assert cov["n_activities"] == 17 and cov["n_constraints"] == 76
    # Every table row maps to a roster activity; every bullet is modeled.
    assert all(r["activity"] for r in cov["rows"])
    assert [b["status"] for b in cov["bullets"]] == ["modeled"] * 10
    assert cov["self_check"]["unchecked"] == []


def test_artemis_roster(artemis):
    acts = {a["id"]: a for a in artemis["scenario"]["activities"]}
    assert len(acts) == 17
    # Rule-stated durations win over the varying table values (warned).
    assert acts["sleep"]["duration"] == 495
    assert any("'Sleep' duration differs" in w for w in artemis["warnings"])
    # Day-varying blocks split per day.
    assert acts["adaptation_d1"]["duration"] == 90 and acts["adaptation_d1"]["days"] == [0]
    assert acts["adaptation_d2"]["duration"] == 120
    assert acts["work_c_d3"]["days"] == [2]
    # Window phrases became daily windows.
    assert acts["adaptation_d1"]["daily_window"] == {"open": "00:00", "close": "12:00"}
    assert acts["ops_sync"]["daily_window"] == {"open": "08:00", "close": "12:00"}
    # Daily activities recur with no pinned time_window.
    assert acts["breakfast"]["recurs_daily"] and acts["breakfast"]["days"] == "all"


def test_artemis_constraints(artemis):
    cons = artemis["scenario"]["constraints"]
    by_type = {}
    for c in cons:
        by_type[c["type"]] = by_type.get(c["type"], 0) + 1
    assert by_type == {"time_lag": 6, "min_separation": 69, "no_overlap": 1}
    # No pinned time_windows: the tables are the oracle, not the extraction.
    assert "time_window" not in by_type

    lags = [c for c in cons if c["type"] == "time_lag"]
    adj = next(c for c in lags if c["from_id"] == "pre_sleep")
    assert (adj["to_id"], adj["min_lag"], adj["max_lag"], adj["day_shift"]) == ("sleep", 0, 0, 0)
    wake = next(c for c in lags if c["to_id"] == "post_sleep")
    assert (wake["from_id"], wake["day_shift"], adj["min_lag"]) == ("sleep", 1, 0)
    # Awake caps: soft target (P3) + hard absolute max (P1), wake -> next sleep.
    caps = sorted((c for c in lags if c["from_id"] == "sleep" and c["to_id"] == "sleep"),
                  key=lambda c: c["max_lag"])
    assert [(c["max_lag"], c["priority"]) for c in caps] == [(945, 3), (990, 1)]
    # Meal chain: start-to-start <= 6h along the day's meal order.
    chain = {(c["from_id"], c["to_id"]) for c in lags if c["max_lag"] == 360}
    assert chain == {("breakfast", "lunch"), ("lunch", "dinner")}
    for c in lags:
        assert c["source"]  # provenance on everything

    seps = [c for c in cons if c["type"] == "min_separation"]
    meals = {frozenset((c["a"], c["b"])) for c in seps if c["gap"] == 30}
    assert meals == {frozenset(("exercise", m)) for m in ("breakfast", "lunch", "dinner")}
    buffer = [c for c in seps if c["gap"] == 10]
    assert len(buffer) == 66
    for c in buffer:  # meals + sleep excluded: their spacing has its own rules
        assert not {c["a"], c["b"]} & {"breakfast", "lunch", "dinner", "sleep"}


def test_artemis_solves_optimal(artemis):
    out = solve(Scenario.model_validate(artemis["scenario"]))
    assert out["status"] == "OPTIMAL"
    assert out["horizon"] == 3 * 1440


def test_artemis_self_check_catches_the_docs_own_inconsistency(artemis):
    # The document contradicts itself once: day-2 sleep (22:00 + 8h30m = 06:30 day 3)
    # overlaps day-3 post-sleep (06:20). The self-check must catch EXACTLY that —
    # everything else in the timetable satisfies the extracted rules.
    v = artemis["coverage"]["self_check"]["violations"]
    assert len(v) == 2
    wake = next(x for x in v if "after waking" in x["label"])
    assert wake["day"] == 1 and "-10m" in wake["detail"]
    ovl = next(x for x in v if "overlap" in x["label"].lower())
    assert ovl["day"] == 2 and "sleep" in ovl["detail"] and "10m" in ovl["detail"]
    assert any("may contradict itself" in w for w in artemis["warnings"])


def test_artemis_doc_timetable_independent_spot_checks(artemis_blocks):
    # Guard the checker itself: recompute a few known facts straight from the tables.
    rows, n_tables, n_days = xs.parse_day_tables(artemis_blocks)
    assert (n_tables, n_days, len(rows)) == (3, 3, 39)
    day2 = {r["name"]: r for r in rows if r["day"] == 1}
    assert day2["Lunch"]["start"] - day2["Breakfast"]["start"] == 335  # <= 360 rule holds
    assert day2["Dinner"]["start"] - day2["Lunch"]["start"] == 360    # exactly at the cap
    assert day2["Sleep"]["end"] < day2["Sleep"]["start"]              # crosses midnight
    assert day2["Exercise"]["start"] - (day2["Lunch"]["end"]) == 30   # >= 30m from meals


def test_extract_route_schedule_genre(client):
    r = client.post("/extract", data={"document": (io.BytesIO(ARTEMIS.read_bytes()),
                                                   "artemis_3day_schedule.docx")})
    assert r.status_code == 200
    data = r.get_json()
    assert data["coverage"]["genre"] == "schedule"
    assert data["scenario"]["horizon"] == 4320
    assert Scenario.model_validate(data["scenario"])
