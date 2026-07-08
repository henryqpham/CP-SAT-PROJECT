"""ingest.extract_blocks on small in-memory .docx files built with python-docx."""
import io
import zipfile

import docx
import pytest

from ingest import extract_blocks


def docx_bytes(fill):
    # build a .docx in memory; fill(doc) adds the content
    doc = docx.Document()
    fill(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def blocks_of(fill):
    return extract_blocks(docx_bytes(fill))["blocks"]


def test_real_heading_style_sets_breadcrumb():
    def fill(doc):
        doc.add_heading("Vehicle Spec", level=1)  # real "Heading 1" style
        doc.add_paragraph("Plain intro text.")
    blocks = blocks_of(fill)
    assert blocks[0]["kind"] == "heading"
    assert blocks[0]["section_path"] == ["Vehicle Spec"]
    assert blocks[1]["kind"] == "text"
    assert blocks[1]["section_path"] == ["Vehicle Spec"]


def test_fake_bold_short_line_is_a_heading():
    def fill(doc):
        para = doc.add_paragraph()  # bold short line, no heading style
        para.add_run("Fake Bold Heading").bold = True
        doc.add_paragraph("Under the fake heading.")
    blocks = blocks_of(fill)
    assert blocks[0]["kind"] == "heading"
    assert blocks[1]["section_path"] == ["Fake Bold Heading"]


def test_bold_divider_line_is_not_a_heading():
    # Docs fake horizontal rules with bold "-----" paragraphs. A punctuation-only
    # line must not become a heading, or it wipes the breadcrumb and shows up as
    # a literal divider in the review table (the testdoc.docx bug).
    def fill(doc):
        para = doc.add_paragraph()
        para.add_run("APPENDIX E — EXPANSION SET").bold = True
        rule = doc.add_paragraph()
        rule.add_run("-" * 60).bold = True
        doc.add_paragraph("[SH-1101] Modeled Synchronization Pulse")
    blocks = blocks_of(fill)
    assert blocks[0]["kind"] == "heading"
    assert blocks[1]["kind"] == "text"  # the divider, demoted
    # the requirement still sits under the appendix title, not the divider
    assert blocks[2]["section_path"] == ["APPENDIX E — EXPANSION SET"]


def test_numbered_paragraph_depth_builds_breadcrumb():
    def fill(doc):
        doc.add_paragraph("4 Braking")
        doc.add_paragraph("4.2 Wheels")
        doc.add_paragraph("4.2.1 Rims")
        doc.add_paragraph("Deep text.")
        doc.add_paragraph("4.3 Axles")
        doc.add_paragraph("Sibling text.")
    blocks = blocks_of(fill)
    assert [b["kind"] for b in blocks[:3]] == ["heading", "heading", "heading"]
    assert blocks[3]["section_path"] == ["4 Braking", "4.2 Wheels", "4.2.1 Rims"]
    # a new depth-2 heading replaces its level and drops anything deeper
    assert blocks[5]["section_path"] == ["4 Braking", "4.3 Axles"]


def test_paragraph_kinds():
    def fill(doc):
        doc.add_paragraph("[VR-110] Brake pads")
        doc.add_paragraph("The system shall stop within 10 m.")
        doc.add_paragraph("Just some plain commentary.")
    blocks = blocks_of(fill)
    assert blocks[0]["kind"] == "requirement"
    assert blocks[0]["requirement_ids"] == ["VR-110"]
    assert blocks[1]["kind"] == "shall"
    assert blocks[1]["is_shall"] is True
    assert blocks[2]["kind"] == "text"


def test_table_cells_become_table_blocks():
    def fill(doc):
        doc.add_heading("Milestones", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "VR-110 due"
        table.cell(0, 1).text = "It shall be done"
    blocks = blocks_of(fill)
    cells = [b for b in blocks if b["kind"] == "table"]
    assert [b["text"] for b in cells] == ["VR-110 due", "It shall be done"]
    assert all(b["section_path"] == ["Milestones"] for b in cells)
    assert cells[0]["requirement_ids"] == ["VR-110"]
    assert cells[1]["is_shall"] is True


def test_all_three_date_formats_found():
    def fill(doc):
        doc.add_paragraph("Due 2026-03-15 and March 16, 2026 and 17 March 2026.")
    (block,) = blocks_of(fill)
    assert block["dates"] == ["2026-03-15", "2026-03-16", "2026-03-17"]


def test_impossible_calendar_date_is_skipped():
    def fill(doc):
        doc.add_paragraph("Ends February 30, 2026 maybe.")
    (block,) = blocks_of(fill)
    assert block["dates"] == []


def test_coverage_counts_and_first_seen_ids():
    def fill(doc):
        doc.add_heading("Spec", level=1)
        doc.add_paragraph("[VR-110] Brake pads")
        doc.add_paragraph("[VR-2] depends on VR-1")
        doc.add_paragraph("The system shall stop.")
        doc.add_paragraph("Mentions VR-110 again.")
    coverage = extract_blocks(docx_bytes(fill))["coverage"]
    assert coverage["requirement_ids"] == ["VR-110", "VR-2", "VR-1"]  # deduped, first-seen order
    assert coverage["n_blocks"] == 5
    assert coverage["n_requirements"] == 2
    assert coverage["n_shall"] == 1


def test_zip_bomb_is_rejected():
    # 81 MB of zeros compresses to almost nothing but claims > the 80 MB cap
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/document.xml", b"\x00" * (81 * 1024 * 1024))
    with pytest.raises(ValueError, match="refusing to parse"):
        extract_blocks(buf.getvalue())


def test_non_zip_bytes_are_rejected():
    # A clean ValueError, so the /extract route can answer 400 (not a 500).
    with pytest.raises(ValueError, match="not a valid .docx"):
        extract_blocks(b"this is not a docx")


def test_accepts_bytes_and_file_like():
    data = docx_bytes(lambda doc: doc.add_paragraph("Hello there."))
    assert extract_blocks(data) == extract_blocks(io.BytesIO(data))
