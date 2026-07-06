"""Track-2 spec extraction: working_window / section_budget / overlap statements,
Rationale lines, and a real priority spread on mixed must/should/may docs.

Small in-memory spec docs exercise each new pattern; the committed sample doc
pins that the new rules add NOTHING there (still 35 constraints) while its
Rationale lines now ride on every derived constraint.
"""
import io

import pytest
from docx import Document

import extract_det as det
from conftest import fail_ask
from ingest import extract_blocks
from extract import extract_document


def spec_blocks(paragraphs):
    """A minimal spec .docx in memory: one numbered heading + the given paragraphs."""
    doc = Document()
    doc.add_paragraph("1 Requirements")
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return extract_blocks(buf.getvalue())["blocks"]


def empty_ask(prompt):
    return {"tasks": []}  # a residual call that fills nothing (falls back to defaults)


DYNO = [
    "[VR-100] Dyno Base Test",
    "The Vehicle System shall complete the base dynamometer cycle.",
    "Rationale: The base cycle anchors all comparative runs.",
    "Estimated validation effort: 1 day. Conducted on the dynamometer. "
    "Testing shall be performed daily between 8:00 and 17:30.",
]


def test_working_window_from_stated_hours():
    res = extract_document(spec_blocks(DYNO), ask=fail_ask)
    ww = [c for c in res["scenario"]["constraints"] if c["type"] == "working_window"]
    assert len(ww) == 1
    w = ww[0]
    assert (w["section"], w["open"], w["close"]) == ("dynamometer", "08:00", "17:30")
    assert w["priority"] == 1  # "shall be performed"
    assert w["rationale"].startswith("The base cycle")
    assert "between 8:00 and 17:30" in w["source"]
    assert res["coverage"]["extraction"]["working_windows"] == 1


def test_working_window_without_resource_is_skipped_not_plan_wide():
    # Hours with no resource must NOT fall back to section="all" (that would turn one
    # requirement's hours into a plan-wide curfew) — skip and warn instead.
    res = extract_document(spec_blocks([
        "[VR-100] Curfew Note",
        "The Vehicle System shall log all activity.",
        "Estimated validation effort: 1 day.",
        "Testing shall be performed daily between 8:00 and 17:00.",
    ]), ask=empty_ask)
    assert not [c for c in res["scenario"]["constraints"] if c["type"] == "working_window"]
    assert res["coverage"]["extraction"]["working_windows"] == 0
    assert any("names no resource" in w for w in res["warnings"])


def test_bare_clock_pair_without_context_is_not_a_window():
    blocks = spec_blocks([
        "[VR-100] Timing Note",
        "The Vehicle System shall log events.",
        "Estimated validation effort: 1 day. Owner: Data Team.",
        "Historic incidents clustered between 2:00 and 3:00.",  # no operations word near it
    ])
    # "clustered" is not an operations word, but "between ... and ..." alone must not fire.
    res = extract_document(blocks, ask=fail_ask)
    assert res["coverage"]["extraction"]["working_windows"] == 0


def test_section_budget_needs_total_and_a_resource():
    res = extract_document(spec_blocks([
        "[VR-100] Dyno Base Test",
        "The Vehicle System shall complete the base dynamometer cycle.",
        "Estimated validation effort: 1 day. Conducted on the dynamometer.",
        "The total dynamometer time shall not exceed 2 days.",
        "[VR-101] Solo Cap",
        "The Vehicle System shall keep the soak test short.",
        # a per-activity cap, NOT a budget: no "total"
        "Estimated validation effort: 4 hours. Owner: Thermal Team. "
        "The soak test shall not exceed 6 hours.",
    ]), ask=fail_ask)
    sb = [c for c in res["scenario"]["constraints"] if c["type"] == "section_budget"]
    assert len(sb) == 1
    assert (sb[0]["section"], sb[0]["max_minutes"]) == ("dynamometer", 2880)
    assert res["coverage"]["extraction"]["section_budgets"] == 1


def test_section_budget_without_resource_is_warned_not_guessed():
    res = extract_document(spec_blocks([
        "[VR-100] Fleet Cap",
        "The Vehicle System shall respect shared lab limits.",
        "Estimated validation effort: 1 day. Owner: Lab Team.",
        # names an aggregate cap but the req's resource is just the owning team —
        # wait, Owner IS a resource; use a req with none at all
    ]), ask=empty_ask)
    # Owner resolves to a resource, so this budget-less doc just has no budget.
    assert res["coverage"]["extraction"]["section_budgets"] == 0

    res = extract_document(spec_blocks([
        "[VR-100] Fleet Cap",
        "The Vehicle System shall respect shared lab limits.",
        "Estimated validation effort: 1 day.",
        "The total test time shall not exceed 1 week.",
    ]), ask=empty_ask)
    assert res["coverage"]["extraction"]["section_budgets"] == 0
    assert any("names no resource" in w for w in res["warnings"])


def test_overlap_during_and_parallel():
    res = extract_document(spec_blocks([
        "[VR-100] Dyno Base Test",
        "The Vehicle System shall complete the base dynamometer cycle.",
        "Estimated validation effort: 1 day. Conducted on the dynamometer.",
        "[VR-101] Telemetry Capture",
        "Capture should be conducted during [VR-100] for the full cycle.",
        "The Vehicle System shall record all channels.",
        "Estimated validation effort: 4 hours. Owner: Data Team.",
        "[VR-102] Noise Survey",
        "The survey may run in parallel with [VR-100] whenever rig capacity allows.",
        "The Vehicle System shall log cabin noise.",
        "Estimated validation effort: 4 hours. Owner: NVH Team.",
    ]), ask=fail_ask)
    ov = {c["inner"]: c for c in res["scenario"]["constraints"] if c["type"] == "overlap"}
    assert set(ov) == {"vr_101", "vr_102"}
    assert ov["vr_101"]["outer"] == "vr_100" and ov["vr_101"]["mode"] == "contains"
    assert ov["vr_101"]["priority"] == 3   # "should be conducted during"
    assert ov["vr_102"]["mode"] == "overlaps"
    assert ov["vr_102"]["priority"] == 5   # "may run in parallel"
    # a real spread across one document: hard shall deps stay 1, should 3, may 5
    assert {c["priority"] for c in res["scenario"]["constraints"]} >= {3, 5}


def test_overlap_unknown_reference_is_skipped_with_warning():
    res = extract_document(spec_blocks([
        "[VR-100] Dyno Base Test",
        "Runs during [VR-999] every cycle.",
        "The Vehicle System shall complete the cycle.",
        "Estimated validation effort: 1 day. Owner: Test Team.",
    ]), ask=fail_ask)
    assert not [c for c in res["scenario"]["constraints"] if c["type"] == "overlap"]
    assert any("VR-999" in w and "skipped" in w for w in res["warnings"])
    assert res["coverage"]["dangling_references"] == []


def test_rationale_rides_on_derived_constraints():
    res = extract_document(spec_blocks([
        "[VR-100] Dyno Base Test",
        "The Vehicle System shall complete the base dynamometer cycle.",
        "Estimated validation effort: 1 day. Conducted on the dynamometer.",
        "[VR-101] Follow-up Test",
        "Rationale: Comparisons only mean something against the base cycle.",
        "The Vehicle System shall repeat the cycle and depends on [VR-100].",
        "Estimated validation effort: 1 day. Conducted on the dynamometer.",
    ]), ask=fail_ask)
    pre = next(c for c in res["scenario"]["constraints"] if c["type"] == "precedence")
    assert pre["rationale"].startswith("Comparisons only mean something")
    assert res["coverage"]["extraction"]["rationales"] == 1


def test_sample_doc_regression_with_rationales(extracted_sample):
    # The new patterns add NOTHING on the sample spec (no stated hours/budgets/overlaps),
    # and every derived constraint now carries its requirement's Rationale line.
    cons = extracted_sample["scenario"]["constraints"]
    assert len(cons) == 35
    assert all(c["rationale"] for c in cons)
    ext = extracted_sample["coverage"]["extraction"]
    assert (ext["working_windows"], ext["section_budgets"], ext["overlaps"]) == (0, 0, 0)
    assert ext["rationales"] == 29
