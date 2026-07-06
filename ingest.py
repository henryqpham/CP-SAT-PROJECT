"""Local .docx -> structured blocks extractor with provenance.

Turns a controlled-document-style spec (numbered sections, [VR-xxx] requirements,
'shall' statements, dates) into an ordered list of blocks, each tagged with the
section breadcrumb it sits under. Everything runs in memory from a file-like
object — nothing is written to disk and nothing leaves the machine (privacy is
non-negotiable, per CLAUDE.md). The returned dict is the contract later phases
build on, so its shape is fixed; see extract_blocks below.
"""
import datetime
import io
import re
import zipfile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# A .docx is a zip of XML; reject one whose entries decompress past this cap BEFORE
# python-docx loads it all into memory (a tiny file can otherwise inflate to GBs — a
# zip bomb). file_size comes from the zip's central directory, so the check is cheap.
MAX_UNCOMPRESSED_BYTES = 80 * 1024 * 1024  # ~80 MB of decompressed XML is plenty for any real spec

# A requirement *id* anywhere in the text (e.g. "depends on VR-110"): two+ caps,
# a dash, digits. A requirement *header* is the same id in brackets at the start.
_REQ_ID = re.compile(r"\b([A-Z]{2,}-\d+)\b")
_REQ_HEADER = re.compile(r"\[([A-Z]{2,}-\d+)\]")

# A faked heading: a dotted section number followed by a word, e.g. "4.2.1 Braking".
# The dotted-number count gives the breadcrumb depth.
_NUMBERED = re.compile(r"^(\d+(?:\.\d+)*)\s+\S")

# Dates. The sample doc uses "Month DD, YYYY"; we also accept ISO and "DD Month YYYY"
# so other specs parse too. Month names are mapped locally — no dateutil dependency.
_MONTHS = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11,
    "december": 12,
}
_ISO_DATE = re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b")
_MONTH_NAMES = "|".join(_MONTHS)
_MDY_DATE = re.compile(rf"\b({_MONTH_NAMES})\s+(\d{{1,2}}),\s+(\d{{4}})\b", re.IGNORECASE)
_DMY_DATE = re.compile(rf"\b(\d{{1,2}})\s+({_MONTH_NAMES})\s+(\d{{4}})\b", re.IGNORECASE)


def _find_dates(text):
    """Return ISO 'YYYY-MM-DD' strings for every date we recognize. Unparseable
    or out-of-range dates are skipped rather than raising."""
    found = []
    for y, m, d in _ISO_DATE.findall(text):
        found.append(_iso(int(y), int(m), int(d)))
    for month, d, y in _MDY_DATE.findall(text):
        found.append(_iso(int(y), _MONTHS[month.lower()], int(d)))
    for d, month, y in _DMY_DATE.findall(text):
        found.append(_iso(int(y), _MONTHS[month.lower()], int(d)))
    return [iso for iso in found if iso]


def _iso(year, month, day):
    """Format a calendar-valid ISO date, or "" (e.g. "February 30" is rejected)."""
    try:
        return datetime.date(year, month, day).isoformat()
    except ValueError:
        return ""


def _is_bold(paragraph):
    """True if the paragraph is entirely (non-empty) bold runs — the sample doc
    fakes headings and requirement labels with bold text, not heading styles."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _heading_level(paragraph):
    """If `paragraph` is a heading, return its 1-based depth; else None.

    Detection is robust to docs that don't use real Word heading styles (the
    sample fakes them with bold numbered paragraphs):
      (a) a "Heading N" style or a non-empty outline level -> use that level;
      (b) a numbered-section pattern ("4.2.1 ...") -> depth = count of numbers;
      (c) a short, bold, Title-ish line -> treat as a top-level heading.
    """
    text = paragraph.text.strip()
    if not text:
        return None

    # (b) Numbered section wins when present: its dotted depth is the truth.
    m = _NUMBERED.match(text)
    if m:
        return m.group(1).count(".") + 1

    # (a) A real Word heading style or outline level.
    style = (paragraph.style.name or "") if paragraph.style else ""
    if style.startswith("Heading"):
        # "Heading 2" -> level 2; bare "Heading" -> level 1.
        tail = style[len("Heading"):].strip()
        return int(tail) if tail.isdigit() else 1
    try:
        outline = paragraph.paragraph_format.outline_level  # None if unset
    except (AttributeError, ValueError):
        outline = None
    if outline is not None:
        return outline + 1  # outline levels are 0-based

    # (c) Short, bold, Title-ish line with no requirement id (e.g. the doc title).
    if _is_bold(paragraph) and len(text) < 80 and not _REQ_HEADER.search(text):
        return 1

    return None


def _classify(text, is_heading):
    """Pick a block kind and collect ids/dates/shall for a piece of text."""
    req_ids = _REQ_ID.findall(text)
    dates = _find_dates(text)
    is_shall = bool(re.search(r"\bshall\b", text, re.IGNORECASE))

    if is_heading:
        kind = "heading"
    elif _REQ_HEADER.search(text):
        # "[VR-110] ..." starts a requirement; this is the header that names it.
        kind = "requirement"
    elif is_shall:
        kind = "shall"
    else:
        kind = "text"
    return kind, req_ids, dates, is_shall


def _block(index, kind, section_path, text, req_ids, dates, is_shall, **extra):
    # `extra` carries optional structure hints: is_bullet on paragraphs,
    # table/row/col on table cells (so a schedule table can be rebuilt).
    return {
        "index": index,
        "kind": kind,
        "section_path": list(section_path),
        "text": text,
        "requirement_ids": req_ids,
        "dates": dates,
        "is_shall": is_shall,
        **extra,
    }


def _guard_zip_bomb(data: bytes):
    """Reject a .docx whose entries decompress beyond MAX_UNCOMPRESSED_BYTES. Reads the
    uncompressed sizes from the zip's central directory (no decompression), so a 300 KB
    bomb that inflates to gigabytes is refused cheaply. A non-zip slips through here and
    is rejected by Document() (the route turns that into a clean 400)."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            total = sum(zi.file_size for zi in zf.infolist())
    except zipfile.BadZipFile:
        return
    if total > MAX_UNCOMPRESSED_BYTES:
        raise ValueError(
            f"document decompresses to ~{total // (1024 * 1024)} MB "
            f"(limit {MAX_UNCOMPRESSED_BYTES // (1024 * 1024)} MB) — refusing to parse"
        )


def extract_blocks(file_like):
    """Extract ordered, provenance-tagged blocks from a .docx file-like object.

    `file_like` is anything python-docx's Document() accepts (an in-memory stream
    or an open file) — we never touch disk. Returns:
        {"blocks": [ {index, kind, section_path, text,
                      requirement_ids, dates, is_shall,
                      is_bullet (paragraphs), table/row/col (table cells)}, ... ],
         "coverage": {requirement_ids, sections, n_blocks,
                      n_requirements, n_dates, n_shall}}
    """
    data = file_like.read() if hasattr(file_like, "read") else file_like
    if isinstance(data, str):
        data = data.encode()
    _guard_zip_bomb(data)
    try:
        doc = Document(io.BytesIO(data))  # parse from memory (caller's stream may be transient)
    except Exception as e:
        # Not a real .docx (garbage bytes, a renamed PDF, a zip that isn't a Word file).
        # Raise ValueError so the /extract route can answer with a clean 400.
        raise ValueError(f"not a valid .docx file ({type(e).__name__})")

    blocks = []
    section_path = []          # running heading breadcrumb, e.g. ["4 ...", "4.2 ..."]
    all_req_ids = []           # distinct requirement ids, in first-seen order
    seen_req_ids = set()
    sections = []              # distinct heading texts, in first-seen order
    seen_sections = set()      # O(1) dedup for `sections` (avoids an O(n^2) scan)
    n_requirements = n_dates = n_shall = 0

    def note_ids(req_ids):
        for rid in req_ids:
            if rid not in seen_req_ids:
                seen_req_ids.add(rid)
                all_req_ids.append(rid)

    # Walk the body in document order so paragraphs and tables stay interleaved.
    # Reading <w:p>/<w:tbl> off the body element skips headers/footers (page-number
    # noise lives there, not in the body).
    table_index = 0  # 0-based, in document order — rides on each cell block
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]  # strip the {namespace}

        if tag == "p":
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            level = _heading_level(para)
            kind, req_ids, dates, is_shall = _classify(text, level is not None)

            if level is not None:
                # A level-N heading replaces the breadcrumb at depth N and drops
                # anything deeper (we're now under a new branch of the tree).
                del section_path[level - 1:]
                section_path.append(text)
                if text not in seen_sections:
                    seen_sections.add(text)
                    sections.append(text)

            # A bulleted/numbered list line ("List Bullet" style) is usually a rule
            # statement in a schedule doc; tag it so extractors can find them.
            style = (para.style.name or "") if para.style else ""
            blocks.append(_block(len(blocks), kind, section_path, text,
                                 req_ids, dates, is_shall,
                                 is_bullet=style.startswith("List")))
            note_ids(req_ids)
            if kind == "requirement":
                n_requirements += 1
            n_dates += len(dates)
            if is_shall:
                n_shall += 1

        elif tag == "tbl":
            # One block per non-empty cell, joining the cell's paragraphs. Cells
            # carry the section breadcrumb in effect, but never change it. table/
            # row/col let an extractor rebuild the table (e.g. a day schedule).
            table = Table(child, doc)
            for r_i, row in enumerate(table.rows):
                for c_i, cell in enumerate(row.cells):
                    text = "\n".join(p.text for p in cell.paragraphs).strip()
                    if not text:
                        continue
                    req_ids = _REQ_ID.findall(text)
                    dates = _find_dates(text)
                    is_shall = bool(re.search(r"\bshall\b", text, re.IGNORECASE))
                    blocks.append(_block(len(blocks), "table", section_path, text,
                                         req_ids, dates, is_shall,
                                         table=table_index, row=r_i, col=c_i))
                    note_ids(req_ids)
                    n_dates += len(dates)
                    if is_shall:
                        n_shall += 1
            table_index += 1

    return {
        "blocks": blocks,
        "coverage": {
            "requirement_ids": all_req_ids,
            "sections": sections,
            "n_blocks": len(blocks),
            "n_requirements": n_requirements,
            "n_dates": n_dates,
            "n_shall": n_shall,
        },
    }
