"""Generate a SYNTHETIC automotive requirements .docx for testing ingestion.

Original content — not derived from any real document. Produces a ~15-page
controlled-document-style spec (header/footer, numbered sections, [VR-xxx]
requirements with bold 'shall', italic Rationale/Effectivity, lettered
sub-lists) — realistic input for the .docx ingestion pipeline.

Beyond the requirement text, each requirement carries SCHEDULING SIGNAL in
natural prose so a downstream parser/LLM can build a multi-day project plan:
  - DURATION/EFFORT  ("Estimated validation effort: 3 days.")
  - DEPENDENCIES     ("shall not begin until [VR-110] is complete", "depends on [VR-xxx]", "after [VR-xxx]")
  - RESOURCES/TEAMS  ("Owner: Chassis Team.", "Requires the shared HIL test bench.")
  - DATED MILESTONES (a "5 Program Milestones" section, mixed date formats)

The content is DETERMINISTIC and literal — no randomness, no external calls.

Run:  python testdata/make_sample_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "sample_vehicle_requirements.docx"

doc = Document()

# --- repeating page header (controlled-doc style) ---
hdr = doc.sections[0].header
hdr.paragraphs[0].text = "Revision: B          |          Document No.: AVP-4400"
hdr.add_paragraph("Release Date: March 4, 2025          |          Page 1 of 15")
hdr.add_paragraph("Title: Vehicle Operational Requirements Specification")

# --- repeating page footer ---
ftr = doc.sections[0].footer
fp = ftr.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    "Proprietary and Confidential — For Internal Engineering Use Only. "
    "Uncontrolled when printed."
)
fr.italic = True


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def major(num, text):
    p = doc.add_paragraph()
    p.add_run(f"{num}  {text.upper()}").bold = True


def sub(num, text):
    p = doc.add_paragraph()
    p.add_run(f"{num}  {text}").bold = True


def body(text):
    doc.add_paragraph(text)


def req(rid, name):
    p = doc.add_paragraph()
    p.add_run(f"[{rid}] {name}").bold = True


def shall(text):
    # bold the word 'shall', like a formal requirements doc
    p = doc.add_paragraph()
    if " shall " in text:
        before, after = text.split(" shall ", 1)
        p.add_run(before + " ")
        p.add_run("shall").bold = True
        p.add_run(" " + after)
    else:
        p.add_run(text)


def italic(prefix, text):
    p = doc.add_paragraph()
    p.add_run(f"{prefix}: {text}").italic = True


def items(letters_and_text):
    for letter, text in letters_and_text:
        p = doc.add_paragraph()
        p.add_run(f"   {letter}. {text}").italic = True


def plan(text):
    # scheduling signal: effort / dependencies / owning resource, in plain prose
    italic("Plan", text)


title("Vehicle Operational Requirements Specification")

body(
    "This specification defines synthetic operational requirements for a battery-electric "
    "passenger vehicle program. Each requirement carries a validation plan stating its "
    "estimated effort, owning resource, and dependencies so that program planning tools can "
    "derive a verification schedule directly from this document. Program milestone dates are "
    "collected in Section 5."
)

# =====================================================================
# 4.2  VEHICLE DYNAMICS AND BRAKING  (original — preserved)
# =====================================================================
major("4.2", "Vehicle Dynamics and Braking")
body(
    "This section defines requirements for vehicle dynamics and braking functions. "
    "Requirements apply to all trim levels unless otherwise stated."
)

sub("4.2.1", "Braking Performance")

req("VR-110", "Antilock Braking Activation")
shall(
    "The Vehicle System shall modulate individual wheel brake pressure to prevent "
    "wheel lockup during braking events on all road surfaces."
)
italic(
    "Rationale",
    "Wheel lockup increases stopping distance and removes steering authority. "
    "Pressure modulation preserves directional control during emergency braking.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 4 days. Owner: Chassis Team. Conducted on the proving ground. "
    "This is a foundational braking item with no prerequisites."
)

req("VR-111", "Regenerative Braking Coordination")
shall(
    "The Vehicle System shall blend regenerative and friction braking to deliver a "
    "consistent deceleration response across the full state-of-charge range."
)
italic(
    "Rationale",
    "Uncoordinated blending produces inconsistent pedal feel and can reduce energy "
    "recovery. Coordinated blending preserves driver expectation while maximizing "
    "recovery. The function shall:",
)
items(
    [
        ("a", "Prioritize regenerative torque when the battery can accept charge;"),
        ("b", "Transition to friction braking smoothly as regenerative limits are reached;"),
        ("c", "Log all blend transitions for diagnostic review."),
    ]
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 3 days. Owner: Chassis Team. Requires the proving ground. "
    "Validation of [VR-111] shall not begin until [VR-110] is complete, because blended "
    "braking is measured against the baseline antilock response."
)

sub("4.2.2", "Stability Control")

req("VR-112", "Electronic Stability Control")
shall(
    "The Vehicle System shall apply selective wheel braking and torque reduction to "
    "correct understeer and oversteer detected from yaw-rate and steering-angle signals."
)
italic(
    "Rationale",
    "Loss of lateral stability is a primary cause of loss-of-control incidents. Selective "
    "intervention restores the intended path with minimal driver workload.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 5 days. Owner: Chassis Team. Conducted on the proving ground. "
    "This activity depends on [VR-110] and shall be performed after [VR-111]."
)

# =====================================================================
# 4.3  DRIVER ASSISTANCE AND MONITORING  (original — preserved + extended)
# =====================================================================
major("4.3", "Driver Assistance and Monitoring")
body(
    "This section defines requirements for driver-assistance and occupant-monitoring "
    "functions."
)

sub("4.3.1", "Driver Attention Monitoring")

req("VR-210", "Driver Drowsiness Detection")
shall(
    "The Vehicle System shall detect driver drowsiness using camera-based gaze and "
    "blink analysis and issue an escalating audible and visual alert upon detection."
)
italic(
    "Rationale",
    "Drowsiness is a leading contributor to single-vehicle incidents. Early detection "
    "gives the driver an opportunity to take corrective action before performance "
    "degrades to an unsafe level.",
)
italic("Effectivity", "Model Year 2027 and later, vehicles with the Driver Monitoring Camera.")
plan(
    "Estimated validation effort: 3 days. Owner: Software Team. Requires the shared HIL test bench."
)

req("VR-211", "Lane-Keeping Assistance")
shall(
    "The Vehicle System shall apply corrective steering torque to maintain the vehicle "
    "within detected lane markings at speeds above 60 km/h."
)
italic(
    "Rationale",
    "Lane departure is a common cause of run-off-road events. Corrective torque reduces "
    "departure frequency while keeping the driver in control.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 4 days. Owner: Software Team. Conducted on the proving ground. "
    "Validation shall not begin until [VR-112] is complete, since lane-keeping torque is "
    "layered on top of stability control."
)

sub("4.3.2", "Collision Avoidance")

req("VR-212", "Forward Collision Warning")
shall(
    "The Vehicle System shall warn the driver of an imminent forward collision at least "
    "two seconds before the projected point of impact."
)
italic(
    "Rationale",
    "A timely warning recovers reaction time that would otherwise be lost, allowing the "
    "driver or automatic braking to mitigate or avoid the event.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 2 days. Owner: Software Team. Requires the shared HIL test bench. "
    "This activity depends on [VR-810], the perception object list."
)

req("VR-213", "Automatic Emergency Braking")
shall(
    "The Vehicle System shall apply automatic braking when a forward collision is imminent "
    "and the driver has not initiated braking."
)
italic(
    "Rationale",
    "When the driver does not respond to a warning, autonomous braking is the last "
    "opportunity to reduce impact energy.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 5 days. Owner: Software Team. Conducted on the proving ground. "
    "Validation of [VR-213] shall not begin until [VR-212] is complete and depends on [VR-110]."
)

# =====================================================================
# 4.4  BATTERY AND THERMAL MANAGEMENT  (original — preserved + extended)
# =====================================================================
major("4.4", "Battery and Thermal Management")
body("This section defines requirements for high-voltage battery thermal protection.")

sub("4.4.1", "Thermal Limits")

req("VR-310", "Battery Over-Temperature Protection")
shall(
    "The Vehicle System shall limit charge and discharge power when any battery cell "
    "exceeds its rated temperature, and shall notify the driver of the limitation."
)
italic(
    "Rationale",
    "Operating cells beyond their rated temperature accelerates degradation and raises "
    "thermal-runaway risk. Power limiting keeps cells within safe bounds while "
    "informing the driver of reduced performance.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 6 days. Owner: Battery Team. Conducted on the environmental "
    "chamber. [VR-310] shall be validated before the 2026-08-15 design freeze."
)

req("VR-311", "Cold-Weather Battery Conditioning")
shall(
    "The Vehicle System shall pre-condition the high-voltage battery to its optimal "
    "charging temperature when a fast-charging session is scheduled in cold ambient conditions."
)
italic(
    "Rationale",
    "Charging a cold battery is slow and accelerates lithium plating. Pre-conditioning "
    "restores charge acceptance and protects long-term capacity.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 8 days. Owner: Battery Team. Conducted on the environmental "
    "chamber. This activity depends on [VR-310] and runs after [VR-310]."
)

sub("4.4.2", "Thermal Runaway Containment")

req("VR-312", "Cell Isolation on Fault")
shall(
    "The Vehicle System shall electrically isolate a battery module within 100 milliseconds "
    "of detecting a cell venting event."
)
italic(
    "Rationale",
    "Rapid isolation limits propagation of a thermal event to neighboring modules and "
    "preserves time for occupant egress.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 7 days. Owner: Battery Team. Requires the battery lab. "
    "Validation shall not begin until [VR-310] is complete."
)

# =====================================================================
# 4.5  POWERTRAIN AND CHARGING  (new)
# =====================================================================
major("4.5", "Powertrain and Charging")
body(
    "This section defines requirements for traction-drive control and the onboard and "
    "off-board charging interfaces."
)

sub("4.5.1", "Traction Drive")

req("VR-410", "Torque Command Arbitration")
shall(
    "The Vehicle System shall arbitrate torque requests from the accelerator, regenerative "
    "braking, and stability functions into a single bounded traction-motor command."
)
italic(
    "Rationale",
    "Multiple sources competing for the same actuator can produce conflicting or unsafe "
    "torque. A single arbitrated command guarantees a predictable response.",
)
italic("Effectivity", "Model Year 2027 and later.")
# cross-section dependency: powertrain (4.5) <- braking (4.2)
plan(
    "Estimated validation effort: 4 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Validation of [VR-410] shall not begin until [VR-111] is complete, because torque "
    "arbitration must respect the regenerative-braking blend."
)

req("VR-411", "Drive Inverter Fault Response")
shall(
    "The Vehicle System shall transition the traction inverter to a safe torque-off state "
    "upon detecting an overcurrent or gate-driver fault."
)
italic(
    "Rationale",
    "An uncontrolled inverter fault can produce unintended torque. A defined safe state "
    "removes drive torque while keeping the vehicle steerable.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "This activity shall take approximately 24 hours. Owner: Software Team. Requires the shared "
    "HIL test bench. Runs after [VR-410]."
)

sub("4.5.2", "Charging Interface")

req("VR-412", "DC Fast-Charge Negotiation")
shall(
    "The Vehicle System shall negotiate the maximum safe charging current with a DC "
    "fast-charging station based on present battery temperature and state of charge."
)
italic(
    "Rationale",
    "Negotiating a current the battery cannot accept causes thermal stress and tripped "
    "sessions. Dynamic negotiation maximizes charging speed within thermal limits.",
)
italic("Effectivity", "Model Year 2027 and later.")
# cross-section dependency: charging (4.5) <- thermal conditioning (4.4)
plan(
    "Estimated validation effort: 5 days. Owner: Battery Team. Conducted on the environmental "
    "chamber. Depends on [VR-311], since negotiation relies on cold-weather conditioning."
)

req("VR-413", "Charge Port Lock Integrity")
shall(
    "The Vehicle System shall keep the charge connector locked while current is flowing and "
    "shall release the lock only after current has stopped."
)
italic(
    "Rationale",
    "Disconnecting under load risks arcing and connector damage. Interlocking the latch to "
    "current flow protects both the vehicle and the operator.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "This activity shall take approximately 12 hours. Owner: Hardware Team. Requires the battery lab. "
    "Runs after [VR-412]."
)

# =====================================================================
# 4.6  CONNECTIVITY AND OVER-THE-AIR UPDATE  (new)
# =====================================================================
major("4.6", "Connectivity and Over-the-Air Update")
body(
    "This section defines requirements for the telematics connection and the over-the-air "
    "(OTA) software update process."
)

sub("4.6.1", "Update Delivery")

req("VR-510", "OTA Package Authentication")
shall(
    "The Vehicle System shall verify the cryptographic signature of every over-the-air "
    "update package before installation and shall reject any package that fails verification."
)
italic(
    "Rationale",
    "An unauthenticated package is an attack vector for arbitrary code execution. Signature "
    "verification ensures only approved software is installed.",
)
italic("Effectivity", "Model Year 2027 and later.")
# cross-section dependency: OTA (4.6) <- cybersecurity (4.9)
plan(
    "Estimated validation effort: 3 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Depends on [VR-910], the secure-boot trust anchor."
)

req("VR-511", "Update Rollback on Failure")
shall(
    "The Vehicle System shall restore the previous software version if an over-the-air "
    "update fails to start the vehicle successfully after installation."
)
italic(
    "Rationale",
    "A failed update must never leave the vehicle inoperable. Automatic rollback preserves "
    "a known-good configuration.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 4 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Validation shall not begin until [VR-510] is complete."
)

sub("4.6.2", "Telematics")

req("VR-512", "Emergency Call Connectivity")
shall(
    "The Vehicle System shall place an automatic emergency call with location data when a "
    "qualifying crash event is detected."
)
italic(
    "Rationale",
    "Automatic notification reduces response time after a severe crash, particularly when "
    "occupants are incapacitated.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "This activity shall take approximately 16 hours. Owner: Software Team. Requires the shared "
    "HIL test bench. Depends on [VR-512] for self-test setup."  # see planted-conflict note below
)
# --- PLANTED INFEASIBILITY #1 (self-loop): [VR-512] depends on itself. A
# dependency resolver that treats this as a precedence edge cannot schedule it
# (a task cannot start after it finishes). This exercises the cycle / self-loop
# detector in the downstream scheduler.

# =====================================================================
# 4.7  HUMAN-MACHINE INTERFACE AND INFOTAINMENT  (new)
# =====================================================================
major("4.7", "Human-Machine Interface and Infotainment")
body(
    "This section defines requirements for the driver display, audible warnings, and the "
    "infotainment system."
)

sub("4.7.1", "Driver Display")

req("VR-610", "Critical Warning Priority")
shall(
    "The Vehicle System shall display safety-critical warnings with priority over all "
    "infotainment content on the primary driver display."
)
italic(
    "Rationale",
    "A critical warning obscured by entertainment content can be missed. Priority "
    "rendering guarantees the driver sees safety messages first.",
)
italic("Effectivity", "Model Year 2027 and later.")
# cross-section dependency: HMI (4.7) <- driver assistance (4.3)
plan(
    "Estimated validation effort: 2 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Depends on [VR-212], whose forward-collision warning is the highest-priority message."
)

req("VR-611", "Display Startup Time")
shall(
    "The Vehicle System shall render the speedometer and critical telltales within two "
    "seconds of an ignition-on event."
)
italic(
    "Rationale",
    "Mandatory telltales must be visible before the vehicle can be driven. A bounded "
    "startup time guarantees the driver has required information.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "This activity shall take approximately 8 hours. Owner: Software Team. Requires the shared "
    "HIL test bench. Runs after [VR-610]."
)

sub("4.7.2", "Audio Warnings")

req("VR-612", "Audible Alert Audibility")
shall(
    "The Vehicle System shall produce safety-critical audible alerts at a level that "
    "remains audible above maximum media playback volume."
)
italic(
    "Rationale",
    "An alert masked by media is ineffective. Guaranteed audibility ensures the driver "
    "perceives the warning regardless of media settings.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 2 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Depends on [VR-610]."
)

# =====================================================================
# 4.8  ADAS PERCEPTION  (new)
# =====================================================================
major("4.8", "ADAS Perception")
body(
    "This section defines requirements for the perception stack that fuses camera, radar, "
    "and ultrasonic sensors into an object model."
)

sub("4.8.1", "Sensor Fusion")

req("VR-810", "Object List Generation")
shall(
    "The Vehicle System shall fuse camera, radar, and ultrasonic data into a unified object "
    "list updated at least every 50 milliseconds."
)
italic(
    "Rationale",
    "Downstream functions act on a single coherent view of the environment. A fused list "
    "reduces conflicting detections and missed objects.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 6 days. Owner: Software Team. Requires the shared HIL test bench. "
    "This is a foundational perception item with no prerequisites."
)

req("VR-811", "Sensor Blockage Detection")
shall(
    "The Vehicle System shall detect when a perception sensor is obstructed and shall "
    "degrade dependent functions gracefully while notifying the driver."
)
italic(
    "Rationale",
    "A blocked sensor silently producing stale data is dangerous. Explicit blockage "
    "detection lets the system reduce capability rather than act on bad data.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 3 days. Owner: Software Team. Requires the shared HIL test bench. "
    "Validation shall not begin until [VR-810] is complete."
)

sub("4.8.2", "Free-Space Estimation")

req("VR-812", "Drivable Path Estimation")
shall(
    "The Vehicle System shall estimate the drivable free space ahead of the vehicle for use "
    "by lane-keeping and emergency-braking functions."
)
italic(
    "Rationale",
    "A reliable free-space estimate bounds where automated functions may steer or brake, "
    "improving robustness in cluttered environments.",
)
italic("Effectivity", "Model Year 2027 and later.")
# cross-section dependency: perception (4.8) <- dynamics (4.2)
plan(
    "Estimated validation effort: 4 days. Owner: Software Team. Conducted on the proving ground. "
    "Depends on [VR-810] and shall be performed after [VR-112], the stability-control baseline."
)

# =====================================================================
# 4.9  CYBERSECURITY  (new)
# =====================================================================
major("4.9", "Cybersecurity")
body(
    "This section defines requirements for vehicle cybersecurity, covering secure boot, "
    "network segmentation, and intrusion detection."
)

sub("4.9.1", "Platform Integrity")

req("VR-910", "Secure Boot Trust Anchor")
shall(
    "The Vehicle System shall validate each boot stage against an immutable hardware trust "
    "anchor before transferring control to that stage."
)
italic(
    "Rationale",
    "A compromised early boot stage undermines all later security controls. A hardware root "
    "of trust ensures only authentic firmware executes.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 5 days. Owner: Security Team. Requires the shared HIL test bench. "
    "This is a foundational security item with no prerequisites."
)

req("VR-911", "In-Vehicle Network Segmentation")
shall(
    "The Vehicle System shall isolate safety-critical control networks from infotainment "
    "networks using an enforcing gateway."
)
italic(
    "Rationale",
    "A flat network lets a compromised infotainment unit reach the brakes or steering. "
    "Segmentation contains an intrusion to its originating domain.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "This activity shall take approximately 32 hours. Owner: Security Team. Requires the shared "
    "HIL test bench. Validation shall not begin until [VR-910] is complete."
)

sub("4.9.2", "Intrusion Detection")

req("VR-912", "Anomaly Detection and Logging")
shall(
    "The Vehicle System shall detect anomalous network traffic on the control bus and shall "
    "log the event for later forensic analysis."
)
italic(
    "Rationale",
    "Detecting and recording an intrusion attempt enables response and improves future "
    "defenses, even when an attack is not fully blocked.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 3 days. Owner: Security Team. Requires the shared HIL test bench. "
    "Depends on [VR-911]."
)

# =====================================================================
# 4.10  MANUFACTURING VALIDATION  (new)
# =====================================================================
major("4.10", "Manufacturing Validation")
body(
    "This section defines end-of-line and durability validation activities performed before "
    "production sign-off."
)

sub("4.10.1", "End-of-Line Test")

req("VR-1010", "End-of-Line Functional Test")
shall(
    "The Vehicle System shall pass a complete end-of-line functional test exercising every "
    "safety-critical function before the vehicle leaves the assembly line."
)
italic(
    "Rationale",
    "A defect that escapes the line is expensive and unsafe to correct in the field. A "
    "comprehensive end-of-line test confirms correct integration of all functions.",
)
italic("Effectivity", "Model Year 2027 and later.")
# This activity aggregates the validated functions, so it gates on the late items.
plan(
    "Estimated validation effort: 5 days. Owner: Manufacturing Team. Conducted on the proving ground. "
    "Validation of [VR-1010] shall not begin until [VR-213] is complete and depends on [VR-812]."
)

req("VR-1011", "Durability Sign-Off Test")
shall(
    "The Vehicle System shall complete an accelerated durability cycle representing ten years "
    "of service without a safety-critical failure."
)
italic(
    "Rationale",
    "Durability defects appear late in service. An accelerated cycle surfaces them before "
    "production commitment.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 10 days. Owner: Manufacturing Team. Conducted on the environmental "
    "chamber. Runs after [VR-1010]."
)

sub("4.10.2", "Pre-Freeze Hardware Validation")

# --- PLANTED INFEASIBILITY #2 (deadline before predecessor can finish):
# [VR-1012] "shall not begin until [VR-1011] is complete" yet must "be complete
# by 2026-08-15" (the design freeze). [VR-1011] itself depends on [VR-1010],
# which depends on the late items [VR-213]/[VR-812]. That chain (e.g.
# VR-810 -> VR-812 -> VR-1010 -> VR-1011 -> VR-1012) totals well over a month of
# effort, so VR-1012 cannot both follow VR-1011 AND finish by the 2026-08-15
# freeze. The conflict is the pair ([VR-1012], 2026-08-15 design freeze).
req("VR-1012", "Pre-Freeze Hardware Qualification")
shall(
    "The Vehicle System hardware build shall be qualified against the structural and thermal "
    "load cases prior to design freeze."
)
italic(
    "Rationale",
    "Hardware changes after freeze are costly. Qualifying the build beforehand confirms the "
    "design is ready to lock.",
)
italic("Effectivity", "Model Year 2027 and later.")
plan(
    "Estimated validation effort: 6 days. Owner: Hardware Team. Conducted on the environmental "
    "chamber. Validation of [VR-1012] shall not begin until [VR-1011] is complete, and [VR-1012] "
    "shall be complete by 2026-08-15."
)

# =====================================================================
# 5  PROGRAM MILESTONES  (new — dated, mixed formats)
# =====================================================================
major("5", "Program Milestones")
body(
    "This section collects the program milestone dates that constrain the validation "
    "schedule defined above. Dates are given in mixed formats intentionally."
)

sub("5.1", "Key Dates")
shall("The design freeze shall be complete by 2026-08-15.")
shall("Software feature complete shall be reached by September 30, 2026.")
shall("The release candidate build shall be available on 12 October 2026.")
shall("The final validation report is due by November 30, 2026.")
shall("Production sign-off: 15 January 2027.")

sub("5.2", "Requirement-to-Milestone Ties")
body(
    "The following ties bind requirement validation to milestone dates. Where a tie cannot "
    "be satisfied, the program plan must surface the conflict for review."
)
shall("[VR-310] shall be validated before the 2026-08-15 design freeze.")
shall("[VR-910] shall be validated before the September 30, 2026 software feature-complete date.")
shall("[VR-510] shall be complete by November 30, 2026, ahead of the final validation report.")
shall(
    "[VR-1011] shall be complete by 15 January 2027 to support production sign-off, and "
    "depends on [VR-1010]."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT}  ({len(doc.paragraphs)} paragraphs)")
